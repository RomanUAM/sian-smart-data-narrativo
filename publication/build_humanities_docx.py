#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path(__file__).resolve().parent


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        shade(table.rows[0].cells[i], "E6E6E6")
    for row in rows:
        cells = table.add_row().cells
        for i, txt in enumerate(row):
            set_cell_text(cells[i], txt)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "F2F2F2")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    p.add_run("\n" + body).font.size = Pt(9.5)
    doc.add_paragraph()


def setup(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True


def build_es() -> None:
    doc = Document()
    setup(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("El tatuaje como archivo narrativo:\nlectura situada de fuentes digitales mediante un sistema local")
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.bold = True

    doc.add_heading("Resumen", level=1)
    doc.add_paragraph(
        "El tatuaje puede leerse como archivo corporal, marca estética, memoria afectiva, oficio visual y objeto de disputa social. "
        "Este documento presenta un procedimiento local para construir un corpus narrativo trazable sin enviar textos a servicios externos. "
        "El ejemplo conductor es el tópico tatuaje, entendido como campo de sentidos: cuerpo, identidad, oficio, estética, salud, regulación, memoria y estigma. "
        "El sistema separa rubros, años y capas de fuente; depura textos; registra vacíos; organiza indicios narrativos; construye mapas de relaciones; y propone síntesis revisables para lectura cercana."
    )
    doc.add_heading("Problema cultural y metodológico", level=1)
    doc.add_paragraph(
        "Una narrativa social no aparece en una sola clase de documento. Puede emerger en conversaciones cotidianas, circular por la prensa, traducirse a lenguaje institucional y estabilizarse después en estudios académicos. "
        "La pregunta no es qué dice Internet sobre el tatuaje, sino cómo distintas capas discursivas producen, desplazan o disputan sentidos alrededor del tatuaje en un periodo y región definidos."
    )
    doc.add_heading("Fases del procedimiento", level=1)
    add_table(
        doc,
        ["Fase", "Qué se hace", "Por qué se hace"],
        [
            ["Observar el campo", "Reconocer el tatuaje como inscripción corporal, práctica visual y relato social.", "Sitúa el problema antes de convertirlo en búsqueda digital."],
            ["Delimitar", "Definir región, periodo y rubros de sentido: identidad, oficio, salud, memoria, diseño.", "Evita mezclar sentidos distintos bajo una misma palabra."],
            ["Reunir voces", "Consultar por capas: prensa, foros, artículos y reportes.", "Distingue conversación pública, prensa, investigación y documentos institucionales."],
            ["Cuidar el corpus", "Eliminar duplicados, menús, publicidad, ruido y homónimos.", "Impide que el ruido web se vuelva resultado interpretativo."],
            ["Interpretar relaciones", "Observar expresiones, actores, momentos narrativos, fuentes y tonalidad léxica.", "Produce indicios y mapas de sentido para lectura crítica."],
            ["Síntesis crítica", "Proponer nodos relevantes y volver a ejemplos textuales.", "Prepara lectura cercana sin presentar el algoritmo como interpretación final."],
        ],
        [1.35, 2.35, 2.35],
    )
    add_callout(doc, "Figura 1. Itinerario interpretativo", "Situar el campo → reunir voces → cuidar el corpus → leer relaciones. La revisión humana ajusta rubros, exclusiones e interpretación.")

    doc.add_heading("Ejemplo conductor: tatuaje", level=1)
    add_table(
        doc,
        ["Rubro", "Variantes de búsqueda"],
        [
            ["Núcleo", "tatuaje, tatuajes, tattoo, tattoos, arte corporal"],
            ["Oficio e industria", "tatuador, tatuadora, estudio de tatuajes, artista del tatuaje"],
            ["Identidad y memoria", "significado de tatuaje, tatuaje e identidad, tatuaje y memoria, tatuaje religioso"],
            ["Sociedad y trabajo", "tatuaje y juventud, tatuaje y empleo, tatuaje y discriminación, tatuaje y género"],
            ["Salud y regulación", "tatuaje y salud, tintas para tatuaje, infección, regulación sanitaria"],
            ["Estética y diseño", "diseño de tatuaje, tatuaje tradicional, tatuaje mexicano, body art"],
        ],
        [1.55, 4.5],
    )
    doc.add_paragraph("Las exclusiones protegen la pregunta de investigación: cigar, tobacco, robusto o colonoscopic tattooing deben retirarse cuando el interés sea tatuaje corporal como práctica cultural.")
    add_callout(doc, "Figura 2. Diseño de corpus", "Pregunta cultural → rubros → años/región → capas de voz → cuidado del texto → lectura asistida.")

    doc.add_heading("Extracción responsable y evidencia parcial", level=1)
    doc.add_paragraph(
        "La corrida secuencial no busca todos los sinónimos a la vez. Primero consulta el término base del rubro y sólo expande a sinónimos si hace falta. "
        "Así se reduce saturación de índices y se conserva qué término produjo cada hallazgo."
    )
    doc.add_paragraph(
        "La extracción respeta límites de fuente. Antes de intentar texto completo se consulta robots.txt. "
        "Si una página no permite extracción, si es paywall o si sólo entrega título/resumen, el registro se conserva como ok_partial. "
        "Ese metadato sirve como señal trazable, pero no debe contarse como cuerpo textual completo."
    )
    add_table(
        doc,
        ["Estado", "Qué significa", "Uso interpretativo"],
        [
            ["ok", "Texto suficiente y permitido para análisis local.", "Puede alimentar n-gramas, red narrativa y lectura cercana."],
            ["ok_partial", "Título, resumen, RSS o metadato sin cuerpo completo.", "Señal de presencia; no equivale a texto completo."],
            ["too_short", "Texto demasiado breve o heredado.", "Debe revisarse antes de usarlo en análisis textual."],
            ["error/fetch_error", "No fue posible recuperar evidencia usable.", "Cuenta como límite de cobertura."],
        ],
        [1.1, 2.45, 2.5],
    )
    add_callout(doc, "Figura 3. Derechos de fuente y grados de evidencia", "Índice público → URL/metadato → revisión de permiso → texto completo u ok_partial → análisis local. El sistema no evade paywalls, sesiones, CAPTCHAs ni espacios privados.")

    doc.add_heading("Red narrativa y tonalidad léxica", level=1)
    doc.add_paragraph(
        "La red narrativa se usa como mapa. Cada texto conecta actores, ideas, fuente, año, lugar y momentos del relato. "
        "La tonalidad léxica no se interpreta como emoción colectiva: funciona como indicador exploratorio de vocabulario valorativo y siempre debe volver a ejemplos textuales."
    )
    add_callout(doc, "Figura 4. Red narrativa mínima", "Texto conectado con actor, momento del relato, idea, fuente, año y lugar. Las aristas indican relaciones como menciona, publica, contiene, usa o sitúa.")
    add_callout(doc, "Figura 5. Síntesis revisable", "Corpus depurado → mapa de relaciones → nodos relevantes → lectura cercana. El algoritmo sugiere entradas; la interpretación vuelve al corpus.")

    doc.add_heading("Salida reproducible", level=1)
    doc.add_paragraph(
        "La aplicación local exporta evidencia estructurada con registros filtrados, procedencia, limpieza, clasificación de fuente, tonalidad léxica, eventos narrativos, actores, grupos de ideas y mapas de relaciones. "
        "Para una publicación robusta falta consolidar un manifiesto de corrida con fecha, parámetros, versión del código, cuotas, semillas y huella del corpus final."
    )
    doc.add_heading("Límites y sesgos", level=1)
    doc.add_paragraph(
        "El corpus no representa por sí mismo a una sociedad: representa una colección situada de huellas disponibles. "
        "Puede haber sesgo de idioma, clase social, alfabetización digital, centralidad urbana, disponibilidad de fuentes, moderación de plataformas, agenda periodística y sobrerrepresentación sanitaria o institucional. "
        "Los foros, Reddit público, blogs y comunidades abiertas son señales de conversación pública indexable; no equivalen a toda la conversación social. "
        "Las exclusiones también son decisiones interpretativas; por eso deben registrarse y justificarse."
    )
    doc.add_heading("Nota técnica mínima", level=1)
    doc.add_paragraph(
        "La selección de nodos relevantes puede formalizarse como cobertura sobre el mapa de relaciones. Se buscan conjuntos pequeños de nodos que mantengan alta relevancia y bajo daño estructural. "
        "Los detalles algorítmicos quedan como apéndice porque la pregunta principal es interpretativa."
    )
    doc.save(OUT / "corpus_narrativo_tatuaje_humanidades_es.docx")


def build_en() -> None:
    doc = Document()
    setup(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("The Tattoo as Narrative Archive:\na situated reading of digital sources with a local system")
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.bold = True

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "The tattoo can be read as a bodily archive, an aesthetic mark, an affective memory, a visual craft, and an object of social dispute. "
        "This document presents a local procedure for building a traceable narrative corpus without sending texts to external services. "
        "The guiding case is tattoo, treated as a field of meanings: body, identity, craft, aesthetics, health, regulation, memory, and stigma."
    )
    doc.add_heading("Cultural and methodological problem", level=1)
    doc.add_paragraph(
        "A social narrative does not appear in only one kind of document. It may emerge in everyday conversation, circulate through journalism, be translated into institutional language, and later become stabilized in academic research. "
        "The question is not what the Internet says about tattoos, but how different discursive layers produce, move, or dispute meanings around tattooing in a defined period and region."
    )
    doc.add_heading("Interpretive phases", level=1)
    add_table(
        doc,
        ["Phase", "What is done", "Why it matters"],
        [
            ["Situating the field", "Tattooing is recognized as bodily inscription, visual practice, and social narrative.", "The cultural problem is framed before it becomes a digital search."],
            ["Delimiting", "Region, period, and meaning rubrics are defined: identity, craft, health, memory, design.", "Different meanings are not collapsed under one term."],
            ["Gathering voices", "Sources are collected by layer: press, forums, articles, and reports.", "Public conversation, journalism, research, and institutional documents remain distinguishable."],
            ["Caring for the corpus", "Duplicates, menus, advertising, noise, and homonyms are removed.", "Web noise is prevented from becoming interpretive evidence."],
            ["Reading relations", "Expressions, actors, narrative moments, sources, and lexical tonality are observed.", "The system produces clues and maps of meaning for critical reading."],
            ["Critical synthesis", "Relevant nodes are proposed and checked against textual examples.", "The algorithm is not treated as the final interpretation."],
        ],
        [1.35, 2.35, 2.35],
    )
    add_callout(doc, "Figure 1. Interpretive itinerary", "Situate the field → gather voices → care for the corpus → read relations. Human review adjusts rubrics, exclusions, and interpretation.")

    doc.add_heading("Guiding case: tattooing", level=1)
    add_table(
        doc,
        ["Rubric", "Search variants"],
        [
            ["Core terms", "tattoo, tattoos, tatuaje, tatuajes, body art"],
            ["Craft and industry", "tattoo artist, tattoo studio, tatuador, tatuadora, artist of tattooing"],
            ["Identity and memory", "tattoo meaning, tattoo and identity, tattoo and memory, religious tattoo"],
            ["Society and work", "tattoo and youth, tattoo and employment, tattoo discrimination, tattoo and gender"],
            ["Health and regulation", "tattoo health, tattoo inks, infection, sanitary regulation"],
            ["Aesthetics and design", "tattoo design, traditional tattoo, Mexican tattoo, body ink"],
        ],
        [1.55, 4.5],
    )
    doc.add_paragraph("Exclusions protect the research question: cigar, tobacco, robusto, or colonoscopic tattooing should be removed when the interest is tattooing as cultural practice.")
    add_callout(doc, "Figure 2. Corpus design", "Cultural question → rubrics → years/region → voice layers → text care → assisted reading.")

    doc.add_heading("Responsible extraction and partial evidence", level=1)
    doc.add_paragraph(
        "The sequential run does not search all synonyms at once. It starts with the base term of each rubric and expands only when needed. "
        "This reduces saturation of public indexes and preserves which term produced each record."
    )
    doc.add_paragraph(
        "Extraction respects source limits. Before full-text retrieval, robots.txt is checked. "
        "If a page is not allowed, is paywalled, or only exposes a title/summary, the record is stored as ok_partial. "
        "Such metadata is a traceable signal, not equivalent to full text."
    )
    add_table(
        doc,
        ["Status", "Meaning", "Interpretive use"],
        [
            ["ok", "Sufficient and permitted text for local analysis.", "Can feed n-grams, narrative networks, and close reading."],
            ["ok_partial", "Title, summary, RSS, or metadata without full body.", "Presence signal; not equivalent to full text."],
            ["too_short", "Text too brief or inherited.", "Must be reviewed before textual analysis."],
            ["error/fetch_error", "No usable evidence could be retrieved.", "Counts as a coverage limit."],
        ],
        [1.1, 2.45, 2.5],
    )
    add_callout(doc, "Figure 3. Source rights and degrees of evidence", "Public index → URL/metadata → permission check → full text or ok_partial → local analysis. The system does not bypass paywalls, sessions, CAPTCHAs, or private spaces.")

    doc.add_heading("Narrative network and lexical tonality", level=1)
    doc.add_paragraph(
        "The narrative network is a map. Each text connects actors, ideas, source, year, place, and moments of the story. "
        "Lexical tonality is not interpreted as collective emotion: it works as an exploratory indicator of evaluative vocabulary and must always return to textual examples."
    )
    add_callout(doc, "Figure 4. Minimal narrative network", "Text connected with actor, narrative moment, idea, source, year, and place. Edges record relations such as mentions, publishes, contains, uses, or situates.")
    add_callout(doc, "Figure 5. Revisable synthesis", "Clean corpus → map of relations → relevant nodes → close reading. The algorithm suggests entries; interpretation returns to the corpus.")

    doc.add_heading("Reproducible export", level=1)
    doc.add_paragraph(
        "The local application exports structured evidence with filtered records, provenance, cleaning decisions, source classification, lexical tonality, narrative events, actors, idea groups, and relation maps. "
        "For robust publication, a run manifest still needs to consolidate date, parameters, code version, quotas, seeds, and a corpus fingerprint."
    )
    doc.add_heading("Limits and biases", level=1)
    doc.add_paragraph(
        "The corpus does not represent a society by itself: it represents a situated collection of available traces. "
        "It may contain biases of language, social class, digital literacy, urban centrality, source availability, platform moderation, journalistic agenda, and sanitary or institutional overrepresentation. "
        "Forums, public Reddit, blogs, and open communities are signals of indexable public conversation; they are not the whole social conversation. "
        "Exclusions are also interpretive decisions; they must therefore be registered and justified."
    )
    doc.add_heading("Minimal technical note", level=1)
    doc.add_paragraph(
        "The selection of relevant nodes can be formalized as coverage over the map of relations. The goal is to find small node sets that preserve high relevance and low structural damage. "
        "Algorithmic details remain in an appendix because the main question is interpretive."
    )
    doc.save(OUT / "narrative_tattoo_corpus_humanities_en.docx")


if __name__ == "__main__":
    build_es()
    build_en()
