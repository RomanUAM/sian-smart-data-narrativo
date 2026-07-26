from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("modelo_cubridor_narrativo.docx")


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Modelo local de red narrativa ponderada y cubridor nodal")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "Sistema local para extracción, limpieza, validación humana, red compleja y selector nodal multiobjetivo inspirado en SCP"
    )
    run.italic = True
    run.font.size = Pt(10.5)


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def math(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], GRAY)
        set_cell_margins(hdr[idx])
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9.2)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def main() -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc)

    h(doc, "1. Objeto y alcance")
    p(
        doc,
        "El sistema analiza narrativas en corpus documentales sin enviar textos a modelos externos. "
        "Extrae documentos, limpia y normaliza texto, clasifica fuentes, identifica etapas narrativas, "
        "permite validación humana de actores y construye una red compleja ponderada para resolver un "
        "cubridor nodal narrativo.",
    )
    p(
        doc,
        "La salida no debe interpretarse como la totalidad de una narrativa pública, sino como una muestra "
        "auditable y reproducible condicionada por consulta, años, región, fuentes activas, exclusiones y errores de descarga.",
    )

    h(doc, "2. Extracción responsable y estados de evidencia")
    p(
        doc,
        "El sistema consulta índices públicos, RSS públicos, páginas abiertas y URLs semilla. No automatiza sesiones, "
        "no evade CAPTCHAs, no salta paywalls y no extrae espacios privados o cerrados. Antes de intentar HTML completo "
        "se consulta robots.txt; si no hay permiso o el acceso es parcial, se conserva sólo metadato trazable.",
    )
    add_table(
        doc,
        ["Estado", "Definición operativa", "Uso en análisis"],
        [
            ["ok", "Texto suficiente y permitido para análisis local.", "Puede alimentar redes, n-gramas y lectura cercana."],
            ["ok_partial", "Título, resumen, RSS o metadato sin cuerpo completo.", "Señal de presencia; no equivale a texto completo."],
            ["too_short", "Texto insuficiente para análisis pleno.", "Debe auditarse antes de interpretar."],
            ["error/fetch_error", "No se recuperó evidencia usable.", "Se reporta como límite de cobertura."],
        ],
        [1.1, 2.75, 2.65],
    )
    p(
        doc,
        "Esta distinción evita que una referencia indexada, un resumen o una página restringida pese igual que un texto completo. "
        "Para foros y redes públicas, la evidencia es una señal indexable parcial, no una representación de toda la conversación social.",
    )

    h(doc, "3. Preprocesamiento y extracción narrativa")
    bullet(doc, "El análisis léxico usa texto normalizado en minúsculas, sin URL, con acentos homogeneizados y stopwords editables.")
    bullet(doc, "Monogramas, bigramas, trigramas y redes de coocurrencia se calculan después de la limpieza.")
    bullet(doc, "Los actores se extraen como candidatos sobre texto editorial original, pero se filtran etiquetas metodológicas, fuentes, tópicos y ruido.")
    bullet(doc, "Cada documento conserva evento inicial, conflicto, punto de cambio, resolución y consecuencias con oración y marcadores detectados.")

    h(doc, "4. Red narrativa")
    math(doc, "G = (V, A, w, tau, sigma)")
    p(
        doc,
        "V contiene documentos, actores, etapas narrativas, fuentes, años, localizaciones, tipos de fuente y conceptos "
        "lingüísticos. A contiene relaciones documento-actor, documento-etapa, actor-etapa, fuente-etapa, año-etapa, "
        "localización-etapa, fuente-concepto y coocurrencias concepto-concepto.",
    )
    p(
        doc,
        "Se agrega una capa de flujo narrativo: diálogo individual/foros → noticias → derivaciones técnicas o institucionales → "
        "investigación → otras derivaciones. Esta capa no es un peso; sirve para estudiar circulación y retroalimentación entre discursos.",
    )
    p(
        doc,
        "El modo base es neutral: los elementos narrativos no reciben pesos distintos por decreto. "
        "Una relación pesa más sólo si aparece más veces en el corpus filtrado. La ponderación por evidencia "
        "se conserva únicamente como análisis de sensibilidad o como análisis estratificado por tipo de fuente.",
    )
    math(doc, "w_a = sum_{d in D} 1{a aparece en d}")
    math(doc, "w_norm_a = w_a / max_b w_b;   c_norm_v = c_v / max_u c_u;   degw_norm_v = degw_v / max_u degw_u")
    math(doc, "sigma_v = g1*c_norm_v + g2*degw_norm_v + g3*ev_v + g4*val_v,   sum g_i = 1")
    p(
        doc,
        "Antes de los grafos semántico y de conocimiento se componen frases canónicas: si un bigrama o trigama repetido explica "
        "la mayoría de apariciones de sus partes, las partes se marcan como absorbidas y la frase completa se usa como nodo principal. "
        "Las aristas de coocurrencia y documento--concepto se construyen sobre esas unidades canónicas. "
        "La frase debe aparecer al menos dos veces para evitar conceptos accidentales.",
    )
    add_table(
        doc,
        ["Fuente", "Uso en modo base", "Uso en sensibilidad"],
        [
            ["Artículo científico abierto", "Estrato separado; no vale más por defecto.", "Puede ponderarse como evidencia alta si se declara."],
            ["Reporte institucional/industrial", "Estrato separado; no vale más por defecto.", "Puede contrastarse contra prensa/foros."],
            ["Noticia", "Narrativa pública mediada por prensa.", "Comparar contra artículos y foros."],
            ["Foro/discusión pública", "Prácticas, opiniones y conflictos de usuarios.", "Comparar como discurso situado."],
            ["No clasificable", "Auditoría; no debe dominar el análisis.", "Excluir o analizar aparte."],
        ],
        [2.4, 1.1, 3.0],
    )

    h(doc, "5. Selector nodal multiobjetivo inspirado en SCP")
    p(
        doc,
        "El universo no son los documentos sino las aristas objetivo A*. El modelo se inspira en cobertura de conjuntos, "
        "pero no debe llamarse SCP clásico estricto mientras no se exija cobertura total del universo. "
        "El tercer objetivo mide daño estructural: una arista se cuenta como removida si toca al menos un nodo seleccionado.",
    )
    add_table(
        doc,
        ["Símbolo", "Definición"],
        [
            ["x_v ∈ {0,1}", "1 si el nodo candidato v es seleccionado."],
            ["r_a ∈ {0,1}", "1 si la arista a se remueve al retirar el conjunto seleccionado C."],
            ["b_av ∈ {0,1}", "1 si la arista a es incidente al nodo v."],
            ["k", "Máximo de nodos interpretables en la solución."],
            ["e_v", "Elegibilidad del nodo después de exclusiones manuales/metodológicas."],
        ],
        [1.5, 5.0],
    )
    h(doc, "5.1 Restricciones", level=2)
    math(doc, "r_a >= b_av x_v   para todo a in A*, v in B")
    math(doc, "r_a <= sum_{v in B} b_av x_v   para todo a in A*")
    math(doc, "sum_{v in B} x_v <= k;     x_v <= e_v")
    p(doc, "Estas restricciones implementan remoción lógica de aristas, presupuesto interpretativo y exclusión de ruido.")

    h(doc, "5.2 Objetivos", level=2)
    add_table(
        doc,
        ["Objetivo", "Forma normalizada", "Sentido"],
        [
            ["f1", "sum x_v / |B|", "Minimizar número relativo de nodos."],
            ["f2", "sum sigma_v x_v / sum sigma_v", "Maximizar importancia nodal explicada."],
            ["f3", "sum w_a r_a / sum w_a", "Minimizar daño estructural por aristas removidas."],
        ],
        [1.0, 2.3, 3.2],
    )
    math(doc, "g(x,r) = (1 - f1(x), f2(x), 1 - f3(r)) in [0,1]^3")
    p(
        doc,
        "Las soluciones se comparan por dominancia de Pareto. Los métodos se comparan con hipervolumen normalizado aproximado "
        "y con una superficie empírica u1 × u3 → u2 para no confundir una proyección 2D con el frente completo.",
    )

    doc.add_page_break()
    h(doc, "6. Métodos de solución")
    p(
        doc,
        "Todos los métodos se evalúan bajo el mismo protocolo: misma instancia, mismas restricciones, mismo criterio "
        "de factibilidad, mismo presupuesto de evaluaciones de función objetivo y las mismas métricas de calidad de frente.",
    )
    add_table(
        doc,
        ["Método", "Estrategia de búsqueda", "Condiciones comunes de evaluación"],
        [
            ["Barrido glotón ponderado", "Construye soluciones para distintos vectores de peso sobre u1,u2,u3.", "Mismo presupuesto E, mismo criterio de factibilidad y mismas métricas."],
            ["MOEA", "Cruza y muta subconjuntos; conserva archivo de no dominados.", "Mismo presupuesto E, mismo criterio de factibilidad y mismas métricas."],
            ["MOSA", "Recocido multiobjetivo con aceptación por dominancia y temperatura.", "Mismo presupuesto E, mismo criterio de factibilidad y mismas métricas."],
            ["MMC-MO", "Compositores modifican arreglos de nodos guiados por memoria LP-Pareto.", "Mismo presupuesto E, mismo criterio de factibilidad y mismas métricas."],
        ],
        [1.7, 3.0, 1.8],
    )
    p(
        doc,
        "Las métricas comunes son: hipervolumen normalizado aproximado, número de puntos no dominados globales, distancia generacional "
        "inversa (IGD) contra el frente global, spacing como uniformidad local y dispersión/extensión del frente en el espacio u1,u2,u3.",
    )

    h(doc, "7. Adaptación MMC-MO con guía PL")
    p(
        doc,
        "El MMC-MO no resuelve relajaciones lineales en cada iteración. Para cada instancia del problema, resuelve una sola vez "
        "las relajaciones monoobjetivo de u1=1-f1, u2=f2, u3=1-f3 y un punto balanceado. Las soluciones fraccionales se redondean "
        "y forman la guía inicial de la sociedad de compositores.",
    )
    math(doc, "M_0 = { C^LP_u1, C^LP_u2, C^LP_u3, C^LP_bal }")
    p(
        doc,
        "Durante la búsqueda, la guía se actualiza con soluciones no dominadas encontradas por los compositores, sin volver a resolver PL.",
    )
    math(doc, "M_{t+1} = ND(M_t union P_t)")
    p(
        doc,
        "La comparación usa el criterio de factibilidad: entre factibles se aplica Pareto; entre factible e infactible se elige la factible; "
        "entre infactibles se conserva la de menor violación total.",
    )

    h(doc, "8. Salidas y auditoría")
    for item in [
        "JSON/JSONL incremental, más archivos por año para evitar pérdida por cierres o timeouts.",
        "Tablas por año, fuente, tipo de fuente, idioma, región y estado de descarga.",
        "N-gramas, redes semánticas, módulos semánticos Louvain/alternativas locales y grafo de conocimiento.",
        "Actores candidatos con validación humana persistente.",
        "Frentes Pareto, hipervolumen aproximado, superficie empírica y comparación de métodos.",
    ]:
        bullet(doc, item)

    h(doc, "9. Limitaciones publicables")
    p(
        doc,
        "Crossref, OpenAlex, GDELT, Google News RSS y Reddit no son universos completos; son mecanismos de muestreo. "
        "Toda publicación debe reportar consulta, variantes, años, región, fuentes activas, exclusiones, tasa de errores "
        "de descarga, proporción ok/ok_partial y una auditoría manual de calidad textual y actores.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
